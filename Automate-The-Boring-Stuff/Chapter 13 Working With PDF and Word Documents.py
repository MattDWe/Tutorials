"""
Automate The Boring Stuff Chapter 13
Working With PDF and Word Documents
Github: MattDWe
"""

import PyPDF2
#Used for parsing PDF files to find text.

#Reading PDFs
pdfFileObj = open(".\\Chapter 13\\meetingminutes.pdf", "rb") #Read in binary mode
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
print(pdfReader.numPages)
pageObj = pdfReader.getPage(0) #Zero based index for pages
print(pageObj.extractText())

pdfFileObj = open(".\\Chapter 13\\encrypted.pdf", "rb")
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
print(pdfReader.isEncrypted)
try:
    print(pdfReader.getPage(0))
except:
    print("This is encrypted so you can not read it until password is given.")
pdfReader.decrypt("rosebud")
print(pdfReader.numPages)

#Creating PDFs
"""
Unlike text files Python can not write to PDF files. Instead it can copy different pages into new files.
It can also rotate pages, overlay pages, and encrypt pages.
Steps to do this:
1. Open one or more existing PDFs into PdfFileReader obj
2. Create a new PdfFileWriter object
3. Copy pages from the PdfFileReader objects into the PdfFileWriter object
4. Finally use the PdfFileWriter object to write the output PDF.
"""

#PDF
#Combine two PDFs into one
pdf1File = open(".\\Chapter 13\\meetingminutes.pdf", "rb")
pdf2File = open(".\\Chapter 13\\meetingminutes2.pdf", "rb")
pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
pdfWriter = PyPDF2.PdfFileWriter()

for pageNum in range(pdf1Reader.numPages):
    pageObj = pdf1Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

for pageNum in range(pdf2Reader.numPages):
    pageObj = pdf2Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

pdfOutputFile = open(".\\Chapter 13\\combinedminutes.pdf", "wb")
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()
#PyPDF2 cannot add pages to the middle it will only add pages to the end.

#Rotate a page and save it to another PDF
minutesFile = open(".\\Chapter 13\\meetingminutes.pdf", "rb")
pdfReader = PyPDF2.PdfFileReader(minutesFile)
page = pdfReader.getPage(0)
page.rotateClockwise(90)
pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(page)
resultPdfFile = open(".\\Chapter 13\\rotated.pdf", "wb")
pdfWriter.write(resultPdfFile)
resultPdfFile.close()
minutesFile.close()

#Overlaying pages ontop of eachother
minutesFile = open(".\\Chapter 13\\meetingminutes.pdf", "rb")
pdfReader = PyPDF2.PdfFileReader(minutesFile)
minutesFirstPage = pdfReader.getPage(0)
pdfWatermarkReader = PyPDF2.PdfFileReader(open(".\\Chapter 13\\watermark.pdf", "rb"))
minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(minutesFirstPage)

for pageNum in range(1, pdfReader.numPages):
    pageObj = pdfReader.getPage(pageNum)
    pdfWriter.addPage(pageObj)
resultPdfFile = open(".\\Chapter 13\\watermarkedCover.pdf", "wb")
pdfWriter.write(resultPdfFile)
minutesFile.close()
resultPdfFile.close()

#Encrypting PDFs
pdfFile = open(".\\Chapter 13\\meetingminutes.pdf", "rb")
pdfReader = PyPDF2.PdfFileReader(pdfFile)
pdfWriter = PyPDF2.PdfFileWriter()
for pageNum in range(pdfReader.numPages):
    pdfWriter.addPage(pdfReader.getPage(pageNum))

pdfWriter.encrypt("swordfish")
resultPdf = open(".\\Chapter 13\\encryptedminutes.pdf", "wb")
pdfWriter.write(resultPdf)
resultPdf.close()

#Word Documents
#Reading Docs
import docx
doc = docx.Document(".\\Chapter 13\\demo.docx")
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[4].text)

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)

print(getText(".\\Chapter 13\\demo.docx"))

#Writing to Word Documents
doc = docx.Document()
doc.add_paragraph("Hello World!")
doc.save(".\\Chapter 13\\helloworld.docx")

doc = docx.Document()
doc.add_paragraph("Hello World!")
paraObj1 = doc.add_paragraph("This is the second paragraph")
paraObj2 = doc.add_paragraph("This is another paragraph")
paraObj1.add_run(" This text is being added to the other paragraph")
doc.add_paragraph("Hello World!", "Title")
doc.save(".\\Chapter 13\\multipleParagraps.docx")

"""
'Normal'
'Heading5'
'ListBullet'
'ListParagraph'
'BodyText'
'Heading6'
'ListBullet2'
'MacroText'
'BodyText2'
'Heading7'
'ListBullet3'
'NoSpacing'
'BodyText3'
'Heading8'
'ListContinue'
'Quote'
'Caption'
'Heading9'
'ListContinue2'
'Subtitle'
'Heading1'
'IntenseQuote'
'ListContinue3'
'TOCHeading'
'Heading2'
'List'
'ListNumber'
'Title'
'Heading3'
'List2'
'ListNumber2'
'Heading4'
'List3'
'ListNumber3'

bold The text appears in bold.
italic The text appears in italic.
underline The text is underlined.
strike The text appears with strikethrough.
double_strike The text appears with double strikethrough.
all_caps The text appears in capital letters.
small_caps The text appears in capital letters, with lowercase letters two points smaller.
shadow The text appears with a shadow.
outline The text appears outlined rather than solid.
rtl The text is written right-to-left.
imprint The text appears pressed into the page.
emboss The text appears raised off the page in relief.

add_heading("Text", "size")
doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
"""

"""
Practice Questions
1. The file object name from open()
2. Read binary for pdffilereader and write binary for pdffilewriter
3. Open using open() create a reading object and create a page object from that
4. numPages
5. decrypt("swordfish")
6. rotateClockwise() and rotateCounterClockwise() with degrees
7. docs.Document("demo.docx")
8. Paragraph is seperated by new line and run is a run until the next style
9. doc.paragraphs
10. A run object 
11. True means bolded false means not bolded None defaults to the runs style
12. docx.Document()
13. doc.add_paragraph("Hello there!")
14. Intergers are 0,1,2,3,4
"""
